from flask import Flask,request,jsonify,send_file
from flask import Flask,request,jsonify, send_file
from flask_cors import CORS
import vonage
from dotenv import load_dotenv
import openai
import chardet
import numpy as np
import tensorflow as tf
import tensorflow_hub as hub
from sklearn.cluster import KMeans
import requests
from bs4 import BeautifulSoup
import os,re
from requests import get
import urllib.request
from urllib.parse import urlparse, urljoin
from gtts import gTTS
import json
from docx import Document
import io
from langchain.document_loaders import YoutubeLoader
from langchain.llms import OpenAI
from langchain.chains.summarize import load_summarize_chain

load_dotenv()
app = Flask(__name__, static_url_path='/static', static_folder='static')
CORS(app)
openai.api_key = "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"

@app.route('/meet', methods=['POST'])
def meet():
    client = vonage.Client(
        application_id="47868911",
        private_key="./private.key",
    )

    print(request.json)
    response = client.meetings.create_room({
   "display_name": "Meeting with teacher",
   "metadata": "Welcome to the meeting",
   "type": "instant",
   "recording_options": {
      "auto_record": "false",
      "record_only_owner": "false"
   },
   "join_approval_level": "none",
   "initial_join_options": {
      "microphone_state": "on"
   },
   "available_features": {
      "is_recording_available": "true",
      "is_chat_available": "true",
      "is_whiteboard_available": "true",
      "is_locale_switcher_available": "true",
      "is_captions_available": "false"
   },
   "ui_settings": {
      "language": "en"
   }
})
    print(response["_links"]["host_url"]["href"])

    return response["_links"]


@app.route('/upload', methods=['POST'])
def upload_files():
    notes_file = request.files['notes']
    syllabus_file = request.files['syllabus']
    pyqs_file = request.files['pyqs']

    # Save the files to your desired location
    notes_file.save('uploads/notes.pdf')
    syllabus_file.save('uploads/syllabus.pdf')
    pyqs_file.save('uploads/pyqs.pdf')

    return 'Files uploaded successfully'

def extract_questions_from_file(filepath):
    with open(filepath, 'rb') as f:
        result = chardet.detect(f.read())
    encoding = result['encoding']
    with open(filepath, encoding=encoding) as f:
        content = f.read()
        pattern = r'((?:[IVX]+|\([a-z]\))\. .*(?:\n\s+\(\w\)\. .*)*)'
        matches = re.findall(pattern, content)
        questions = [re.sub(r'\n\s+\(\w\)\. ', ' ', match.strip()) for match in matches]
    return questions

def extract_questions_from_directory(directory):
    questions = []
    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)
        if os.path.isfile(filepath):
            questions += extract_questions_from_file(filepath)
    return questions


def cluster_questions(questions, num_clusters, syllabus_file):

    module_url = "https://tfhub.dev/google/universal-sentence-encoder-large/5"

    embed = hub.load(module_url)
    embeddings = embed(questions).numpy()
    kmeans = KMeans(n_clusters=num_clusters)
    kmeans.fit(embeddings)
    print("ok")
    y_kmeans = kmeans.predict(embeddings)

    return y_kmeans


# Process each paragraph and search for related images
def process_paragraphs(paragraphs):
    # Create the "images" directory if it doesn't exist
    if not os.path.exists("images"):
        os.makedirs("images")
    for query in paragraphs:
        image_url = search_image(query)
        if image_url:
            # Download the image and save it locally
            response = requests.get(image_url)
            if response.status_code == 200:
                with open(f"images/image_{query}.jpg", 'wb') as f:
                    f.write(response.content)
                    print(f"Image_{query}.jpg saved.")

def search_image(query):
    url = f"https://www.google.com/search?q={query}&tbm=isch&tbs=isz:l"
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/58.0.3029.110 Safari/537.3"
    }
    response = requests.get(url, headers=headers)
    soup = BeautifulSoup(response.text, 'html.parser')
    image_elements = soup.find_all('img')
    image_url = None
    if len(image_elements) > 1:
        image_url = image_elements[1].get('src')
        if image_url:
            parsed_url = urlparse(image_url)
            if not parsed_url.scheme:
                base_url = response.url
                image_url = urljoin(base_url, image_url)
    return image_url


def extract_text_from_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except FileNotFoundError:
        print(f"File '{file_path}' not found.")
        return None

def extract_explanation(questions):
    text = questions
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "user",
                # "content": f"You are a teacher, who will take a Topic to a student in classroom. I will provide you with a basic summary of the student's lecture note, which may be a bit incomplete,but you have to use it and should be easy to the student and make sure they understand it fully. Basically, you have to narrate in atmost 300 words and in the end of the lecture you have to summarize the entire lecture in 100 words and say that lets summarize our entire lecture and then tell the summary and don't say Thank you or any other greeting in beginning or last of lecture(everything should be text):\n\n{text}\n\n"
                "content": f"You are a teacher, who will take a Topic to a student in classroom. I will provide you with a basic summary of the student's lecture note, which may be a bit incomplete,but you have to use it and should be easy to the student and make sure they understand it fully. Basically, you have to narrate in atmost 300 words and don't say Thank you or any other greeting in beginning or last of lecture(everything should be text):\n\n{text}\n\n"
            }
        ]
    )

    important_topics = response.choices[0].message.content
    print(important_topics)
    return important_topics

def extract_important_topics(questions):
    text = questions
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "user",
                "content": f"create a important topics for each paragraph as a list and give only 10-15 word in list\n\n{text}\n\n"
            }
        ]
    )

    important_topics = response.choices[0].message.content
    print(important_topics)
    return important_topics


def extract_image_name(questions):
    text = questions

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {
                "role": "user",
                "content": f"create a important topics for each paragraph as a list and give only 10 word in list\n\n{text}\n\n"
            }
        ]
    )

    important_topics = response.choices[0].message.content
    print("2")
    print(questions)
    print("3")
    print(important_topics)
    return important_topics


@app.route('/')
def hello():
    return 'Hello, World!'

@app.route('/ai-tutor')
async def tutor():
    print("Giving promt to OPENAI...")
    # file_path = 'Files\generated_files\summarised_notes\\module1_summarized.txt'
    file_path = 'module1_summarized.txt'
    extracted_text = extract_text_from_file(file_path)
    narrate = extract_explanation(extracted_text)

    print("--------narrate----------")
    print(narrate)

    # Convert text to speech
    tts = gTTS(narrate)

    # Save speech to a temporary file
    speech_file = 'output.mp3'
    tts.save(f"./static/{speech_file}")

    # speech_file = 'output.mp3'
    # tts.save(speech_file)

    print("Video output saved")

    # return send_file(speech_file), narrate
    return narrate


@app.route('/get-audio')
async def getAudio():
    return send_file("output.mp3")



@app.route('/imp-questions')
async def impquestion():
    print("Giving promt to OPENAI...")
    file_path = 'Files\generated_files\\narration.txt'
    extracted_text = extract_text_from_file(file_path)
    narrate = extract_important_topics(extracted_text)

    print("--------narrate----------")
    print(narrate)

    return narrate


@app.route('/pyqs-questions')
async def pyqquestion():
    print("Extracting question paper text")
    questions = extract_questions_from_directory('Files/pyqs_text')
    # num_clusters = int(input("To how many clusters do you want to cluster the questions: "))
    num_clusters = 4
    syllabus_file = 'Files/syllabus_txt/syllabus.txt'
    print("Extracting syllabus")
    labels = cluster_questions(questions, num_clusters, syllabus_file)

    print("Clustering questions")
    for i in range(num_clusters):
        cluster_questions1 = np.array(questions)[np.where(labels == i)[0]]
        print(f"Module {i+1}:")
        for question in cluster_questions1:
            print(f" - {question}")


    questionsArr = []

    # Save cluster questions to file
    with open('Files/generated_files/cluster_questions.txt', 'w') as f:
        for i in range(num_clusters):
            cluster_questions1 = np.array(questions)[np.where(labels == i)[0]]
            f.write(f"Module {i+1}:\n")
            for question in cluster_questions1:
                f.write(f" - {question}\n")
                questionsArr.append(question)
            f.write("\n")



    return questionsArr

@app.route('/get-yt', methods=['POST'])
async def youtubevideo():
    question = request.json["ques"]
    print(question)
    OPENAI_API_KEY = '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'

    youtube_search_api = "https://www.googleapis.com/youtube/v3/search"

    params = {
        "part": "snippet",
        "q": question,
        "key": "AIzaSyCAUBhMWXt8VSJbCI9tMo9K452YeoWUWxI",
        "type": "video",
        "maxResults": 1
    }

    response = requests.get(youtube_search_api, params=params)

    if response.status_code == 200:
        data = response.json()
        if 'items' in data and len(data['items']) > 0:
            video_id = data['items'][0]['id']['videoId']
            print(video_id)

            youtube_video_link = f"https://www.youtube.com/watch?v={video_id}"

            loader = YoutubeLoader.from_youtube_url(f"{youtube_video_link}", add_video_info=True)

            result = loader.load()

            print(result)

            print (f"Found video from {result[0].metadata['author']} that is {result[0].metadata['length']} seconds long")

            llm = OpenAI(temperature=0, openai_api_key=OPENAI_API_KEY)
            chain = load_summarize_chain(llm, chain_type="stuff", verbose=False)
            res = chain.run(result)

            return jsonify(youtube_video_link=youtube_video_link,summary=res)
        else:
            return jsonify(error="No YouTube videos found for the given prompt")
    else:
        return jsonify(error="Failed to fetch YouTube video link")

# Store notes in database
@app.route('/store-notes', methods=['POST'])
def store_notes():
   print(request.json["data"])

         # Define the URL of the Flask route
   url = 'https://ap-south-1.aws.neurelo.com/rest/notes/__one'  # Change the URL to match your Flask server address

   # Define the data to be sent in the request
   data = request.json["data"]

   # Define the headers
   headers = {
      'Content-Type': 'application/json',  # Specify the content type as JSON
      'X-API-KEY': 'neurelo_9wKFBp874Z5xFw6ZCfvhXZLGfamMor4qKoFilXKGGpNFi5P4IoL1PYJoZwsJqpSNxXtuRdh852eBD2nS9eVZxeOBlqvzaBbxqfn3HU8Odfw3HVoDRYZhjVwhhBQCQM0Aw133v3B3ld2YcmNAL2FNTAwQscRYSJ+z7XTbKUSxSefZTxqEUDKQ8ZQcHQ9N8RwS_RdMGvZNTD5LaGNqUNU1n9bBqnLtsTGDG/B9fK1FAvmI='
   }

   # Make the POST request using the requests module
   response = requests.post(url, json=data, headers=headers)

   # Check if the request was successful (status code 200)
   if response.status_code == 201:
      # Print the response from the server
      print(response.json())
   else:
      # Print an error message if the request was not successful
      print('Error:', response.status_code)
   return jsonify({'message': 'Notes stored successfully'})

# Get notes from database of a particular user
@app.route('/get-notes', methods=['GET'])
def get_notes():
   # Access query parameters from the URL
   param1 = request.args.get('param1')


   # Define the URL of the Flask route
   base_url = 'https://ap-south-1.aws.neurelo.com/rest/notes'  # Change the URL to match your Flask server address

   # Define the headers
   headers = {
      'X-API-KEY': 'neurelo_9wKFBp874Z5xFw6ZCfvhXZLGfamMor4qKoFilXKGGpNFi5P4IoL1PYJoZwsJqpSNxXtuRdh852eBD2nS9eVZxeOBlqvzaBbxqfn3HU8Odfw3HVoDRYZhjVwhhBQCQM0Aw133v3B3ld2YcmNAL2FNTAwQscRYSJ+z7XTbKUSxSefZTxqEUDKQ8ZQcHQ9N8RwS_RdMGvZNTD5LaGNqUNU1n9bBqnLtsTGDG/B9fK1FAvmI='
   }

   # Make the GET request using the requests module
   response = requests.get(base_url, params=param1, headers=headers)

   # Check if the request was successful (status code 200)
   if response.status_code == 200:
      # Print the response from the server
      # Parse JSON data
      parsed_data = response.json()

      # Create a new Document
      doc = Document()

      # Add content to the Document
      for entry in parsed_data['data']:
         doc.add_heading(f"Topic: {entry['topic']}", level=1)
         doc.add_paragraph(f"Content: {entry['content']}")
         doc.add_paragraph()  # Add an empty paragraph for spacing

      # Save the Document to a file
      doc.save('formatted_notes.docx')

      # Save the Document to a BytesIO object
      output = io.BytesIO()
      doc.save(output)
      output.seek(0)

      # Send the file as a response
      return send_file(output,download_name='formatted_notes.docx', as_attachment=True)

      # return response.json()
   else:
      # Print an error message if the request was not successful
      return {'Error:' : response.status_code }

if __name__ == "__main__":
    app.run(debug=True)